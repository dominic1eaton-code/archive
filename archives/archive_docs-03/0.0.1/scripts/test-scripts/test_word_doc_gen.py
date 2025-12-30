from docx import Document
from docx.shared import Inches
import datetime

documents = {
    'software development contract',
    'articles of organization',
    'articles of association',
    'memorandum of association',
    'software development agreement',
    'sharedholder agreement',
    'general partnership agreement',
    'limited partnership agreement'
}

def test_doc():
    document = Document()
    document.add_paragraph("It was a dark and stormy night.")
    document.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.")

    # <docx.text.paragraph.Paragraph object at 0x10f19e760>
    document.save("dark-and-stormy.docx")

    document = Document("dark-and-stormy.docx")
    document.paragraphs[0].text
    # 'It was a dark and stormy night.'


def generate_software_development_contract():
    sender_company = "Bahari Software Contracting Firm"
    client_company = "Moyo Software Firm"
    firm = "Moyo"
    sender_company_location = "123 west applegate street, st louis MO"
    intro = f"This Software Development Agreement (the \"Agreement\" or \"Software Development Agreement\") states the terms of and conditions that govern the contractual agreement between {sender_company} having their principle place of business at {sender_company_location} \
             (the \"Developer\") and {client_company}"
    section = f"The Developer shall provide software development services to create [Description of Software/Project], hereinafter referred to as the \"Software.\" The Software shall include [List Key Features and Functionalities]."
    section2 = f"Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum."
    note1 = f"revised: {datetime.datetime.now()}"
    doc = [intro, section, section2, note1]
    document = Document()
    document.add_heading(f'Software Develoment Contract \n {client_company}', 0)
    document_name = f"software-development-contract-{firm}.docx"
    for paragraph in doc:
        document.add_paragraph(paragraph)
    document.add_page_break()
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    document.add_picture('linuxlogo.png', width=Inches(1.25))
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    document.add_page_break()
    document.save(document_name)

if __name__ == '__main__':
    generate_software_development_contract()
